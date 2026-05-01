from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from jobhunt.pipeline.tailor import (
    TailoredCategory,
    TailoredResume,
    TailoredRole,
)
from jobhunt.resume.render_docx import estimate_lines, fits_one_page, render


@pytest.fixture
def tailored() -> TailoredResume:
    return TailoredResume(
        summary=(
            "Full-stack developer with hands-on Shopify, HubSpot, and React work. "
            "Owns project lifecycles end-to-end. GBC diploma, Dean's List."
        ),
        skills_categories=[
            TailoredCategory("Core", ["JavaScript", "TypeScript", "React", "Next.js"]),
            TailoredCategory("CMS", ["Shopify (Liquid)", "HubSpot CMS"]),
            TailoredCategory("Familiar", ["Java", "Python"]),
        ],
        roles=[
            TailoredRole(
                title="Web Developer (Contract)",
                employer="Custom Jewelry Brand (NDA)",
                dates="2023 – Present",
                bullets=["Built 14+ page Shopify storefront.", "Shipped ring builder."],
            ),
        ],
        certifications=["Contentful Certified Professional (October 2025)"],
        education=["Computer Programming & Analysis (Advanced Diploma), GBC, April 2024"],
        coursework=["Data Structures & Algorithms", "Full-Stack Development"],
        model="qwen3:14b",
    )


def test_render_writes_valid_docx(tailored: TailoredResume, tmp_path: Path):
    out = render(
        tailored,
        contact_line="me@example.com | site.com",
        name="Casey Hsu",
        out_path=tmp_path / "out.docx",
    )
    assert out.exists()
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Casey Hsu" in text
    assert "SUMMARY" in text
    assert "TECHNICAL SKILLS" in text
    assert "Familiar" in text
    assert "Custom Jewelry Brand" in text


def test_estimate_fits_one_page(tailored: TailoredResume):
    assert fits_one_page(tailored)
    assert estimate_lines(tailored) > 0


def test_render_emits_single_deans_list_paragraph(tailored: TailoredResume, tmp_path: Path):
    out = render(
        tailored,
        contact_line="me@example.com",
        name="Casey Hsu",
        out_path=tmp_path / "out.docx",
    )
    doc = Document(str(out))
    deans = [p for p in doc.paragraphs if p.text.startswith("Dean")]
    assert len(deans) == 1, [p.text for p in deans]
