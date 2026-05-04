"""Render a CoverLetter to an ATS-safe .docx (matches resume styling)."""

from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from jobhunt.pipeline.cover import CoverLetter
from jobhunt.resume.render_docx import (
    BODY_FONT,
    BODY_SIZE,
    NAME_SIZE,
    _scrub_metadata,
    _set_default_font,
    _set_margins,
    _tighten,
)


def _add_letter_paragraph(doc: Any, text: str, *, after: int = 8) -> None:
    """Cover letter body paragraph: full Pt(after) gap, no first-line indent."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = BODY_SIZE
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.15


def render_cover(cover: CoverLetter, contact_line: str, name: str, out_path: Path) -> Path:
    doc = Document()
    _set_margins(doc)
    _set_default_font(doc)
    _scrub_metadata(doc, name)

    p_name = doc.add_paragraph()
    p_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r_name = p_name.add_run(name)
    r_name.bold = True
    r_name.font.name = BODY_FONT
    r_name.font.size = NAME_SIZE
    _tighten(p_name, after=2)

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r_contact = p_contact.add_run(contact_line)
    r_contact.font.name = BODY_FONT
    r_contact.font.size = BODY_SIZE
    _tighten(p_contact, after=18)

    _add_letter_paragraph(doc, date.today().strftime("%B %d, %Y"), after=12)
    _add_letter_paragraph(doc, cover.salutation, after=10)

    for para in cover.body:
        text = para.strip()
        if text:
            _add_letter_paragraph(doc, text, after=10)

    sign_off_lines = [line for line in cover.sign_off.split("\n") if line.strip()]
    for line in sign_off_lines:
        _add_letter_paragraph(doc, line, after=2)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
