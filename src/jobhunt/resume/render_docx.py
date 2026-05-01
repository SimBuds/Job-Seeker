"""Render a TailoredResume to an ATS-safe .docx (Calibri, single-column, real bullets).

ATS rules enforced (Resume_Tailoring_Instructions.md §5):
- Single column. No tables for layout. No graphics, headers, or footers.
- Calibri 10–11pt body, 14–16pt name.
- Real bullet list style (not typed asterisks).
- US Letter, 0.5"–0.75" margins.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from jobhunt.pipeline.tailor import TailoredResume

BODY_FONT = "Calibri"
BODY_SIZE = Pt(10.5)
NAME_SIZE = Pt(16)
HEADING_SIZE = Pt(11)
HEADING_COLOR = RGBColor(0x33, 0x33, 0x33)


def render(tailored: TailoredResume, contact_line: str, name: str, out_path: Path) -> Path:
    doc = Document()
    _set_margins(doc)
    _set_default_font(doc)

    _add_name(doc, name)
    _add_contact(doc, contact_line)

    _add_section_heading(doc, "SUMMARY")
    _add_paragraph(doc, tailored.summary)

    _add_section_heading(doc, "TECHNICAL SKILLS")
    for cat in tailored.skills_categories:
        if not cat.items:
            continue
        p = doc.add_paragraph()
        run_label = p.add_run(f"{cat.name}: ")
        run_label.bold = True
        run_label.font.name = BODY_FONT
        run_label.font.size = BODY_SIZE
        run_items = p.add_run(", ".join(cat.items))
        run_items.font.name = BODY_FONT
        run_items.font.size = BODY_SIZE
        _tighten(p)

    _add_section_heading(doc, "PROFESSIONAL EXPERIENCE")
    for role in tailored.roles:
        p = doc.add_paragraph()
        run_t = p.add_run(f"{role.title} | {role.employer}")
        run_t.bold = True
        run_t.font.name = BODY_FONT
        run_t.font.size = BODY_SIZE
        # Right-align dates via tab + tab-stop at the right margin.
        run_d = p.add_run(f"\t{role.dates}")
        run_d.font.name = BODY_FONT
        run_d.font.size = BODY_SIZE
        _add_right_tab_stop(p)
        _tighten(p)
        for bullet in role.bullets:
            bp = doc.add_paragraph(style="List Bullet")
            br = bp.add_run(bullet)
            br.font.name = BODY_FONT
            br.font.size = BODY_SIZE
            _tighten(bp)

    _add_section_heading(doc, "CERTIFICATIONS & EDUCATION")
    for line in tailored.certifications:
        _add_paragraph(doc, line)
    for line in tailored.education:
        _add_paragraph(doc, line)
    if tailored.coursework:
        p = doc.add_paragraph()
        r1 = p.add_run("Dean's List (all terms). ")
        r1.bold = True
        r1.font.name = BODY_FONT
        r1.font.size = BODY_SIZE
        r2 = p.add_run("Coursework: " + ", ".join(tailored.coursework) + ".")
        r2.font.name = BODY_FONT
        r2.font.size = BODY_SIZE
        _tighten(p)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _set_margins(doc: Any) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def _set_default_font(doc: Any) -> None:
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)


def _add_name(doc: Any, name: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(name)
    r.bold = True
    r.font.name = BODY_FONT
    r.font.size = NAME_SIZE
    _tighten(p, after=2)


def _add_contact(doc: Any, contact_line: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(contact_line)
    r.font.name = BODY_FONT
    r.font.size = BODY_SIZE
    _tighten(p, after=6)


def _add_section_heading(doc: Any, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = BODY_FONT
    r.font.size = HEADING_SIZE
    r.font.color.rgb = HEADING_COLOR
    _tighten(p, before=6, after=2)
    _add_bottom_border(p)


def _add_paragraph(doc: Any, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = BODY_SIZE
    _tighten(p)


def _tighten(paragraph: Any, *, before: int = 0, after: int = 2) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.15


def _add_right_tab_stop(paragraph: Any) -> None:
    from docx.enum.text import WD_TAB_ALIGNMENT

    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)


def _add_bottom_border(paragraph: Any) -> None:
    from docx.oxml import OxmlElement

    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pbdr.append(bottom)
    p_pr.append(pbdr)


# --- one-page heuristic check -------------------------------------------------

# Conservative budget for 10.5pt Calibri on US Letter with 0.5"/0.75" margins.
# The previous 52-line budget overshot — section headings, role headers, and
# skill lines that wrap each cost more than the flat 1 line we counted.
# 48 lines with wrap-aware skill/summary/bullet counts produces reliable
# single-page output.
LINES_PER_PAGE = 48
BULLET_CHARS_PER_LINE = 95
SUMMARY_CHARS_PER_LINE = 100
SKILL_CHARS_PER_LINE = 95


def _wrapped_lines(text: str, width: int) -> int:
    return max(1, (len(text) + width - 1) // width)


def estimate_lines(tailored: TailoredResume) -> int:
    lines = 4  # name + contact + spacing
    lines += 1 + _wrapped_lines(tailored.summary, SUMMARY_CHARS_PER_LINE)
    lines += 1
    for cat in tailored.skills_categories:
        if not cat.items:
            continue
        line_text = f"{cat.name}: " + ", ".join(cat.items)
        lines += _wrapped_lines(line_text, SKILL_CHARS_PER_LINE)
    lines += 1
    for role in tailored.roles:
        lines += 1
        for b in role.bullets:
            lines += _wrapped_lines(b, BULLET_CHARS_PER_LINE)
    lines += 1 + len(tailored.certifications) + len(tailored.education)
    if tailored.coursework:
        lines += 2
    return lines


def fits_one_page(tailored: TailoredResume) -> bool:
    return estimate_lines(tailored) <= LINES_PER_PAGE
